"""
Tests for Currency Invoice DOCX Export Service (TDD - tests first)

Tests for generate_currency_invoice_docx:
1. Returns bytes
2. Returned bytes are valid DOCX (PK magic bytes)
3. DOCX contains the invoice number text
4. DOCX contains seller and buyer company names
5. DOCX contains item data (product name, quantity, price)
6. DOCX contains total amount
"""

from io import BytesIO
import sys
import os

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from services.currency_invoice_export import generate_currency_invoice_docx


# ---------------------------------------------------------------------------
# Shared test fixtures
# ---------------------------------------------------------------------------

def _make_invoice():
    return {
        "invoice_number": "CI-Q202601-0004-EUR-EURTR-1",
        "generated_at": "2026-02-25T12:00:00Z",
        "currency": "EUR",
        "total_amount": 68808.20,
        "segment": "EURTR",
    }


def _make_seller():
    return {
        "name": "EURO INVEST CAPITAL DIS TICARET LTD. STI",
        "address": "Oruchreis mah. Tekstilkent cad. No:12 b, ESENLER/ISTANBUL",
        "tax_id": "TR-9876543",
    }


def _make_buyer():
    return {
        "name": "MASTER BEARING LLC",
        "address": "109428, Moscow, Ryazan Ave., 22, room 2",
        "tax_id": "RU-1234567890",
    }


def _make_items():
    return [
        {
            "product_name": "MGC 10 GR 2 38/0",
            "sku": "MGC-001",
            "manufacturer": "CAMA",
            "quantity": 8790,
            "unit": "kg",
            "hs_code": "8484900000",
            "price": 2.98,
            "total": 26194.20,
        },
        {
            "product_name": "MGC 10 GR 2 38/15",
            "sku": "MGC-002",
            "manufacturer": "CAMA",
            "quantity": 14300,
            "unit": "kg",
            "hs_code": "8484900000",
            "price": 2.98,
            "total": 42614.00,
        },
    ]


# ===========================================================================
# TEST: Returns bytes
# ===========================================================================

class TestGenerateDocxReturnsBytes:
    """generate_currency_invoice_docx should return valid bytes."""

    def test_returns_bytes(self):
        result = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_valid_docx_magic_bytes(self):
        """DOCX is a ZIP archive, starts with PK magic bytes."""
        result = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        assert result[:2] == b"PK"


# ===========================================================================
# TEST: Contains invoice number
# ===========================================================================

class TestDocxContainsInvoiceNumber:
    """DOCX should contain the invoice number in the document text."""

    def test_invoice_number_in_document(self):
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        all_text = _extract_all_text(doc)
        assert "CI-Q202601-0004-EUR-EURTR-1" in all_text


# ===========================================================================
# TEST: Contains seller and buyer company names
# ===========================================================================

class TestDocxContainsCompanyNames:
    """DOCX should contain both seller and buyer company names."""

    def test_seller_name_present(self):
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        all_text = _extract_all_text(doc)
        assert "EURO INVEST CAPITAL DIS TICARET LTD. STI" in all_text

    def test_buyer_name_present(self):
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        all_text = _extract_all_text(doc)
        assert "MASTER BEARING LLC" in all_text

    def test_seller_address_present(self):
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        all_text = _extract_all_text(doc)
        assert "ESENLER/ISTANBUL" in all_text

    def test_buyer_address_present(self):
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        all_text = _extract_all_text(doc)
        assert "Moscow" in all_text


# ===========================================================================
# TEST: Contains item data
# ===========================================================================

class TestDocxContainsItemData:
    """DOCX should contain product names, quantities, and prices in the items table."""

    def test_product_name_in_table(self):
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        table_text = _extract_table_text(doc)
        assert "MGC 10 GR 2 38/0" in table_text

    def test_second_product_name_in_table(self):
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        table_text = _extract_table_text(doc)
        assert "MGC 10 GR 2 38/15" in table_text

    def test_quantity_in_table(self):
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        table_text = _extract_table_text(doc)
        # Quantity 8790 should appear (possibly formatted with spaces/commas)
        assert "8790" in table_text or "8,790" in table_text or "8 790" in table_text

    def test_unit_price_in_table(self):
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        table_text = _extract_table_text(doc)
        assert "2.98" in table_text

    def test_manufacturer_in_table(self):
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        table_text = _extract_table_text(doc)
        assert "CAMA" in table_text

    def test_hs_code_in_table(self):
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        table_text = _extract_table_text(doc)
        assert "8484900000" in table_text


# ===========================================================================
# TEST: Contains total amount
# ===========================================================================

class TestDocxContainsTotalAmount:
    """DOCX should contain the total amount."""

    def test_total_amount_in_document(self):
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        all_text = _extract_all_text(doc)
        # Total 68808.20 should appear somewhere (table total row or paragraph)
        assert "68808.20" in all_text or "68,808.20" in all_text or "68 808.20" in all_text


# ===========================================================================
# TEST: Edge cases
# ===========================================================================

class TestDocxEdgeCases:
    """Edge cases for DOCX generation."""

    def test_single_item(self):
        """Single item invoice should generate valid DOCX."""
        items = [_make_items()[0]]
        result = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), items
        )
        assert isinstance(result, bytes)
        assert result[:2] == b"PK"

    def test_empty_optional_fields(self):
        """Invoice with missing optional fields should not crash."""
        seller = {"name": "Seller Co", "address": "", "tax_id": ""}
        buyer = {"name": "Buyer Co", "address": "", "tax_id": ""}
        items = [
            {
                "product_name": "Widget",
                "sku": "",
                "manufacturer": "",
                "quantity": 1,
                "unit": "pcs",
                "hs_code": "",
                "price": 10.0,
                "total": 10.0,
            },
        ]
        invoice = {
            "invoice_number": "CI-TEST-1",
            "generated_at": "2026-01-01T00:00:00Z",
            "currency": "USD",
            "total_amount": 10.0,
            "segment": "TRRU",
        }
        result = generate_currency_invoice_docx(invoice, seller, buyer, items)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_date_formatting(self):
        """Date should be formatted as DD.MM.YYYY."""
        from docx import Document

        docx_bytes = generate_currency_invoice_docx(
            _make_invoice(), _make_seller(), _make_buyer(), _make_items()
        )
        doc = Document(BytesIO(docx_bytes))
        all_text = _extract_all_text(doc)
        assert "25.02.2026" in all_text


# ===========================================================================
# Helpers
# ===========================================================================

def _extract_all_text(doc) -> str:
    """Extract all text from paragraphs and tables."""
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _extract_table_text(doc) -> str:
    """Extract text from all tables only."""
    parts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)
