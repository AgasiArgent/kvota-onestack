"""
Currency Invoice DOCX Export Service

Generates DOCX files for internal currency invoices between group companies.
Format is based on the real invoice template (e.g. "Invojs 167 EURO").

Document structure:
  1. Header table: INVOICE title, buyer/supplier info, date, number
  2. Items table: Name, Manufacturer, Qty, HS Code, Unit Price, Total
  3. TOTAL row at bottom of items table
  4. Payment terms paragraphs
  5. Supplier bank details
  6. Signature block
"""

from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def generate_currency_invoice_docx(
    invoice: dict,
    seller: dict,
    buyer: dict,
    items: list[dict],
) -> bytes:
    """Generate a DOCX file for a currency invoice.

    Args:
        invoice: Invoice metadata dict with keys:
            invoice_number, generated_at, currency, total_amount, segment.
        seller: Seller (supplier) company dict with keys:
            name, address, tax_id.
        buyer: Buyer company dict with keys:
            name, address, tax_id.
        items: List of item dicts, each with keys:
            product_name, sku, manufacturer, quantity, unit,
            hs_code, price, total.

    Returns:
        DOCX file as bytes.
    """
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4 width
    section.page_height = Inches(11.69)  # A4 height
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    # -- Format date --
    invoice_date = _format_date(invoice.get("generated_at", ""))
    invoice_number = invoice.get("invoice_number", "")
    currency = invoice.get("currency", "EUR")
    total_amount = invoice.get("total_amount", 0)

    # =================================================================
    # 1. HEADER TABLE (matches real invoice layout)
    # =================================================================
    header_table = doc.add_table(rows=5, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(header_table)

    # Row 0: INVOICE title (merged across all columns)
    _merge_row(header_table, 0)
    cell_title = header_table.cell(0, 0)
    cell_title.text = ""
    p = cell_title.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INVOICE")
    run.bold = True
    run.font.size = Pt(16)

    # Row 1: BUYER | (empty) | SUPPLIER headers
    _set_cell_text(header_table.cell(1, 0), "BUYER", bold=True, size=Pt(10))
    _set_cell_text(header_table.cell(1, 1), "", size=Pt(10))
    _set_cell_text(header_table.cell(1, 2), "SUPPLIER", bold=True, size=Pt(10))

    # Row 2: Buyer details | (empty) | Supplier details
    buyer_text = buyer.get("name", "")
    buyer_addr = buyer.get("address", "")
    if buyer_addr:
        buyer_text += f"\nAddress: {buyer_addr}"

    seller_text = seller.get("name", "")
    seller_addr = seller.get("address", "")
    if seller_addr:
        seller_text += f"\nAddress: {seller_addr}"

    _set_cell_text(header_table.cell(2, 0), buyer_text, size=Pt(9))
    _set_cell_text(header_table.cell(2, 1), "", size=Pt(9))
    _set_cell_text(header_table.cell(2, 2), seller_text, size=Pt(9))

    # Row 3: Invoice Date
    _set_cell_text(header_table.cell(3, 0), "Invoice Date", bold=True, size=Pt(10))
    _set_cell_text(header_table.cell(3, 1), invoice_date, size=Pt(10))
    _set_cell_text(header_table.cell(3, 2), invoice_date, size=Pt(10))

    # Row 4: Invoice Number
    _set_cell_text(header_table.cell(4, 0), "Invoice Number", bold=True, size=Pt(10))
    _set_cell_text(header_table.cell(4, 1), invoice_number, size=Pt(10))
    _set_cell_text(header_table.cell(4, 2), invoice_number, size=Pt(10))

    # -- Spacer --
    doc.add_paragraph("")

    # =================================================================
    # 2. ITEMS TABLE
    # =================================================================
    col_headers = [
        "Name",
        "Manufacturer",
        f"Quantity ({_get_common_unit(items)})",
        "HS Code",
        f"Unit price, {currency}",
        f"Total price, {currency}",
    ]

    items_table = doc.add_table(rows=1, cols=len(col_headers))
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(items_table)

    # Header row
    for i, header in enumerate(col_headers):
        cell = items_table.cell(0, i)
        _set_cell_text(cell, header, bold=True, size=Pt(9))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, "D9E2F3")

    # Data rows
    for item in items:
        row = items_table.add_row()
        row.cells[0].text = str(item.get("product_name", ""))
        row.cells[1].text = str(item.get("manufacturer", ""))
        row.cells[2].text = _format_quantity(item.get("quantity", 0))
        row.cells[3].text = str(item.get("hs_code", ""))
        row.cells[4].text = _format_price(item.get("price", 0))
        row.cells[5].text = _format_price(item.get("total", 0))

        # Right-align numeric columns
        for idx in (2, 4, 5):
            row.cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Set font size for all cells
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"

    # TOTAL row
    total_row = items_table.add_row()
    _set_cell_text(total_row.cells[0], "TOTAL", bold=True, size=Pt(9))
    total_row.cells[1].text = ""
    # Total quantity
    total_qty = sum(float(item.get("quantity", 0)) for item in items)
    _set_cell_text(total_row.cells[2], _format_quantity(total_qty), bold=True, size=Pt(9))
    total_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row.cells[3].text = ""
    total_row.cells[4].text = ""
    _set_cell_text(
        total_row.cells[5],
        f"{_format_price(total_amount)}",
        bold=True,
        size=Pt(9),
    )
    total_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # -- Spacer --
    doc.add_paragraph("")

    # =================================================================
    # 3. PAYMENT TERMS
    # =================================================================
    _add_paragraph(
        doc,
        "The Parties agreed on the following payment procedure according to this Invoice:",
        bold=True,
        size=Pt(9),
    )

    payment_terms = [
        "The payment for the delivery of a batch of Products defined by this Invoice "
        "is made in advance in the amount of at least 20 (twenty)% of the total amount "
        "of delivery according to this Invoice, within 3 (three) banking days from the "
        "date of the invoice received from the Supplier.",
        "The Buyer, at his discretion, has the right to make an advance in a larger amount, "
        "up to 100 (one hundred) % of the total amount of delivery according to this Invoice, "
        "both in a single amount and in payments.",
        "In the event if The Buyer pays an advance payment in the amount of less than 100%, "
        "then the final payment for the delivered batch of the Products is implemented within "
        "5 (five) working days from the date of the invoice by the Supplier.",
        f"The parties agreed that the Invoice currency is {currency}, "
        "the settlement currency is the Russian Ruble (RUB).",
    ]
    for term in payment_terms:
        _add_paragraph(doc, term, size=Pt(9))

    # -- Spacer --
    doc.add_paragraph("")

    # =================================================================
    # 4. SUPPLIER INFO & BANK DETAILS
    # =================================================================
    seller_name = seller.get("name", "")
    seller_addr = seller.get("address", "")
    seller_tax_id = seller.get("tax_id", "")

    _add_paragraph(doc, seller_name, bold=True, size=Pt(10))
    if seller_addr:
        _add_paragraph(doc, f"Address: {seller_addr}", size=Pt(9))
    if seller_tax_id:
        _add_paragraph(doc, f"Tax ID: {seller_tax_id}", size=Pt(9))

    # -- Spacer --
    doc.add_paragraph("")

    # =================================================================
    # 5. SIGNATURE
    # =================================================================
    _add_paragraph(doc, "General Director", size=Pt(10))
    _add_paragraph(doc, "_________________________", size=Pt(10))

    # -- Serialize to bytes --
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ===========================================================================
# Internal helpers
# ===========================================================================

def _format_date(date_str: str) -> str:
    """Format ISO date string to DD.MM.YYYY."""
    if not date_str:
        return datetime.now().strftime("%d.%m.%Y")
    try:
        dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        return dt.strftime("%d.%m.%Y")
    except (ValueError, TypeError):
        return datetime.now().strftime("%d.%m.%Y")


def _format_price(value) -> str:
    """Format a numeric value as a price string (2 decimal places)."""
    try:
        return f"{float(value):.2f}"
    except (ValueError, TypeError):
        return "0.00"


def _format_quantity(value) -> str:
    """Format quantity -- integer if whole number, otherwise 2 decimals."""
    try:
        fval = float(value)
        if fval == int(fval):
            return str(int(fval))
        return f"{fval:.2f}"
    except (ValueError, TypeError):
        return "0"


def _get_common_unit(items: list[dict]) -> str:
    """Get the most common unit from items, default 'pcs'."""
    if not items:
        return "pcs"
    units = [item.get("unit", "pcs") for item in items]
    # Simple majority
    from collections import Counter
    counts = Counter(units)
    return counts.most_common(1)[0][0] if counts else "pcs"


def _set_cell_text(cell, text: str, bold: bool = False, size=None):
    """Set cell text with formatting, clearing existing content."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if size:
        run.font.size = size
    run.font.name = "Arial"


def _set_cell_shading(cell, color_hex: str):
    """Set cell background shading color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading_elm.append(shading)


def _merge_row(table, row_idx: int):
    """Merge all cells in a row."""
    row = table.rows[row_idx]
    first_cell = row.cells[0]
    for cell in row.cells[1:]:
        first_cell.merge(cell)


def _set_table_borders(table):
    """Set thin borders on all cells in a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)

    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.makeelement(
            qn(f"w:{edge}"),
            {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "0",
                qn("w:color"): "000000",
            },
        )
        borders.append(element)
    tbl_pr.append(borders)


def _add_paragraph(doc, text: str, bold: bool = False, size=None):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.bold = bold
    if size:
        run.font.size = size
    run.font.name = "Arial"
    return p
