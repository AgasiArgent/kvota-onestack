"""
Contract-style Specification DOCX Export Service

Generates editable Word documents matching the PDF specification format.
Uses same data source (fetch_contract_spec_data) as PDF export.

Beta feature - allows users to edit the document if needed.
"""

import io
from typing import Dict, Any, List

from docx import Document
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from services.contract_spec_export import (
    fetch_contract_spec_data,
    format_signatory_name,
    DELIVERY_CONDITIONS_TEMPLATE,
    CALC_FIELDS,
    VAT_RATE,
)
from services.export_data_mapper import (
    format_date_russian_long,
    format_number_russian,
    amount_in_words_russian,
    get_currency_symbol,
    qty_in_words,
)


def _set_cell_border(cell, **kwargs):
    """
    Set cell borders. Usage:
    _set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key, value in kwargs[edge].items():
                element.set(qn('w:{}'.format(key)), str(value))
            tcBorders.append(element)

    tcPr.append(tcBorders)


def _set_table_borders(table):
    """Set borders for all cells in table."""
    border_style = {"sz": 4, "val": "single", "color": "000000"}

    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                top=border_style,
                bottom=border_style,
                left=border_style,
                right=border_style
            )


def _add_paragraph(doc, text: str, bold: bool = False, alignment=None,
                   font_size: int = 10, space_after: int = 6):
    """Add a paragraph with standard formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.bold = bold

    if alignment:
        p.alignment = alignment

    p.paragraph_format.space_after = Pt(space_after)

    return p


def _build_delivery_conditions_list(data: Dict[str, Any]) -> List[str]:
    """
    Build delivery conditions list from fixed template + variable values.
    Returns list of strings (one per condition).
    """
    spec = data["specification"]
    quote = data["quote"]
    customer = data["customer"]
    contract = data["contract"]
    calc_vars = data.get("calc_variables", {})

    # Extract variable values
    contract_number = contract.get("contract_number", "б/н")
    contract_date = format_date_russian_long(contract.get("contract_date"))

    # Payment terms
    advance_from_client = calc_vars.get("advance_from_client", 100)
    advance_val = float(advance_from_client)
    payment_percent = int(advance_val * 100) if advance_val <= 1 else int(advance_val)
    payment_days = int(calc_vars.get("time_to_advance", 5))

    # Delivery days
    delivery_days_val = spec.get("delivery_days")
    if delivery_days_val:
        delivery_days = int(delivery_days_val)
    else:
        calc_delivery = calc_vars.get("delivery_time")
        delivery_days = int(calc_delivery) if calc_delivery else 30

    days_type = spec.get("delivery_days_type", "рабочих дней")

    # Customer addresses
    client_name = customer.get("company_name") or customer.get("name") or "Покупатель"
    registration_address = customer.get("address") or "не указан"
    postal_address = customer.get("postal_address")
    delivery_address = postal_address if postal_address else registration_address
    warehouse_address = delivery_address

    # Incoterms
    delivery_terms = quote.get("delivery_terms", "DDP")
    incoterms = f"{delivery_terms} 2020" if delivery_terms else "DDP 2020"

    # Build conditions list
    conditions = []

    # 1. Quality
    conditions.append(DELIVERY_CONDITIONS_TEMPLATE["quality"])

    # 2. Transport
    conditions.append(DELIVERY_CONDITIONS_TEMPLATE["transport"])

    # 3. Payment terms
    payment_text = DELIVERY_CONDITIONS_TEMPLATE["payment"].format(
        contract_number=contract_number,
        contract_date=contract_date,
        payment_percent=payment_percent,
        payment_days=payment_days
    )
    conditions.append(payment_text)

    # 4. Partial delivery
    conditions.append(DELIVERY_CONDITIONS_TEMPLATE["partial"])

    # 5. Responsibility
    conditions.append(DELIVERY_CONDITIONS_TEMPLATE["responsibility"])

    # 6. Warehouse address
    warehouse_text = DELIVERY_CONDITIONS_TEMPLATE["warehouse"].format(
        warehouse_address=warehouse_address
    )
    conditions.append(warehouse_text)

    # 7. Delivery time
    delivery_time_text = DELIVERY_CONDITIONS_TEMPLATE["delivery_time"].format(
        delivery_days=delivery_days,
        days_type=days_type
    )
    conditions.append(delivery_time_text)

    # 8. Consignee (multiline)
    consignee_text = DELIVERY_CONDITIONS_TEMPLATE["consignee"].format(
        client_name=client_name,
        registration_address=registration_address,
        delivery_address=delivery_address
    )
    conditions.append(consignee_text)

    # 9. Incoterms
    incoterms_text = DELIVERY_CONDITIONS_TEMPLATE["incoterms"].format(
        incoterms=incoterms
    )
    conditions.append(incoterms_text)

    return conditions


def generate_contract_spec_docx(spec_id: str, org_id: str) -> bytes:
    """
    Generate DOCX for contract-style specification.

    Args:
        spec_id: Specification UUID
        org_id: Organization UUID

    Returns:
        DOCX file as bytes
    """
    # 1. Fetch data using existing function
    data = fetch_contract_spec_data(spec_id, org_id)

    spec = data["specification"]
    quote = data["quote"]
    items = data["items"]
    customer = data["customer"]
    seller_company = data["seller_company"]
    contract = data["contract"]
    signatory = data["signatory"]
    calculations = data["calculations"]
    organization = data["organization"]
    spec_count = data.get("spec_count", 1)

    # 2. Extract values
    contract_number = contract.get("contract_number", "б/н")
    contract_date = format_date_russian_long(contract.get("contract_date"))
    spec_date = format_date_russian_long(spec.get("specification_date") or spec.get("sign_date"))

    seller_name = seller_company.get("name") or spec.get("our_legal_entity") or organization.get("name", "Поставщик")
    customer_name = customer.get("company_name") or customer.get("name") or spec.get("client_legal_entity", "Покупатель")

    # Director names
    seller_director_parts = [
        seller_company.get("general_director_last_name", ""),
        seller_company.get("general_director_first_name", ""),
        seller_company.get("general_director_patronymic", "")
    ]
    seller_director_full = " ".join(p for p in seller_director_parts if p)
    seller_director = format_signatory_name(seller_director_full)

    customer_signatory_full = signatory.get("name", "")
    customer_signatory = format_signatory_name(customer_signatory_full)
    signatory_position = signatory.get("position", "Генеральный директор")

    # Currency
    spec_currency = spec.get("specification_currency") or quote.get("currency", "RUB")
    currency_symbol = get_currency_symbol(spec_currency)

    # Calculate totals
    if calculations:
        total_no_vat = calculations.get("calc_ae16_sale_price_total") or calculations.get("total_no_vat", 0)
        total_with_vat = calculations.get("calc_al16_total_with_vat") or calculations.get("total_with_vat", 0)

        if total_with_vat:
            totals = {
                "total_qty": sum(item.get("quantity", 1) for item in items),
                "total_no_vat": float(total_no_vat) if total_no_vat else 0,
                "total_with_vat": float(total_with_vat) if total_with_vat else 0,
                "vat_amount": (float(total_with_vat) if total_with_vat else 0) - (float(total_no_vat) if total_no_vat else 0),
            }
        else:
            totals = None
    else:
        totals = None

    # Fallback calculation
    if not totals:
        from decimal import Decimal
        totals = {
            "total_qty": 0,
            "total_no_vat": Decimal("0"),
            "total_with_vat": Decimal("0"),
            "vat_amount": Decimal("0"),
        }
        for item in items:
            calc = item.get("calc", {})
            qty = max(item.get("quantity") or 1, 1)
            totals["total_qty"] += qty
            totals["total_no_vat"] += Decimal(str(calc.get(CALC_FIELDS["TOTAL_NO_VAT"], 0)))
            totals["total_with_vat"] += Decimal(str(calc.get(CALC_FIELDS["TOTAL_WITH_VAT"], 0)))
        totals["vat_amount"] = totals["total_with_vat"] - totals["total_no_vat"]
        totals = {k: float(v) if isinstance(v, Decimal) else v for k, v in totals.items()}

    total_with_vat = totals["total_with_vat"]
    vat_amount = totals["vat_amount"]
    total_qty = totals["total_qty"]

    amount_words = amount_in_words_russian(float(total_with_vat), spec_currency)
    qty_words = qty_in_words(total_qty)

    # Unit form
    last_digit = total_qty % 10
    last_two = total_qty % 100
    if last_two in (11, 12, 13, 14):
        unit_form = "штук (единиц)"
    elif last_digit == 1:
        unit_form = "штука (единица)"
    elif last_digit in (2, 3, 4):
        unit_form = "штуки (единицы)"
    else:
        unit_form = "штук (единиц)"

    # 3. Create Word document
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # 4. Add content sections

    # === Header (Appendix reference) ===
    _add_paragraph(doc, f"Приложение № {spec_count}", font_size=10, space_after=2)
    _add_paragraph(doc, f"к договору поставки № {contract_number}", font_size=10, space_after=2)
    _add_paragraph(doc, f"от {contract_date}", font_size=10, space_after=12)

    # === Title ===
    _add_paragraph(doc, f"СПЕЦИФИКАЦИЯ № {spec_count}", bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14, space_after=6)

    # === Subtitle ===
    _add_paragraph(doc, f"к договору поставки № {contract_number} от {contract_date}",
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=10, space_after=2)
    _add_paragraph(doc, f"между {seller_name} и {customer_name}",
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=10, space_after=12)

    # === Date/Location row ===
    date_table = doc.add_table(rows=1, cols=2)
    date_table.columns[0].width = Cm(8)
    date_table.columns[1].width = Cm(8)

    cell_left = date_table.cell(0, 0)
    cell_left.text = "г. Москва"
    cell_left.paragraphs[0].runs[0].font.size = Pt(10)

    cell_right = date_table.cell(0, 1)
    cell_right.text = spec_date
    cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell_right.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()  # Spacer

    # === Items Table (8 columns) ===
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths (total ~16cm for A4 with margins)
    col_widths = [Cm(0.8), Cm(2.0), Cm(2.0), Cm(4.0), Cm(1.6), Cm(1.2), Cm(2.2), Cm(2.2)]

    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # Header row
    headers = ["№", "IDN-SKU", "Артикул", "Наименование продукции", "Бренд",
               "Кол-во", f"Цена в т.ч. НДС ({VAT_RATE}%)", f"Общая стоимость в т.ч. НДС ({VAT_RATE}%)"]

    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for idx, item in enumerate(items, 1):
        calc = item.get("calc", {})
        qty = max(item.get("quantity") or 1, 1)

        item_total_vat = float(calc.get(CALC_FIELDS["TOTAL_WITH_VAT"], 0))
        price_per_unit_vat = item_total_vat / qty if qty > 0 else 0

        row = table.add_row()
        row_data = [
            str(idx),
            str(item.get('idn_sku') or '-'),
            str(item.get('product_code', '-')),
            str(item.get('product_name', '')),
            str(item.get('brand', '-')),
            str(qty),
            f"{format_number_russian(price_per_unit_vat)} {currency_symbol}",
            f"{format_number_russian(item_total_vat)} {currency_symbol}"
        ]

        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(9)

            # Align numbers to right, center for qty and №
            if i in (0, 5):  # № and Qty
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i in (6, 7):  # Prices
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Totals row
    totals_row = table.add_row()
    totals_row.cells[0].merge(totals_row.cells[4])
    totals_row.cells[0].text = "Итого:"
    totals_row.cells[0].paragraphs[0].runs[0].bold = True
    totals_row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    totals_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    totals_row.cells[5].text = str(total_qty)
    totals_row.cells[5].paragraphs[0].runs[0].bold = True
    totals_row.cells[5].paragraphs[0].runs[0].font.size = Pt(9)
    totals_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    totals_row.cells[6].text = ""

    totals_row.cells[7].text = f"{format_number_russian(total_with_vat)} {currency_symbol}"
    totals_row.cells[7].paragraphs[0].runs[0].bold = True
    totals_row.cells[7].paragraphs[0].runs[0].font.size = Pt(9)
    totals_row.cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Apply borders to table
    _set_table_borders(table)

    doc.add_paragraph()  # Spacer

    # === Summary section ===
    _add_paragraph(doc, "Итог:", bold=True, font_size=10, space_after=4)
    _add_paragraph(doc, f"- общее количество поставляемой Продукции по настоящей Спецификации составляет {total_qty} ({qty_words}) {unit_form} Продукции.",
                   font_size=10, space_after=4)
    _add_paragraph(doc, "- общая сумма поставки Продукции по настоящей Спецификации:",
                   font_size=10, space_after=2)
    _add_paragraph(doc, f"    {format_number_russian(total_with_vat)} ({amount_words}),",
                   bold=True, font_size=10, space_after=2)
    _add_paragraph(doc, f"    в т.ч. НДС {VAT_RATE}% - {format_number_russian(vat_amount)}.",
                   font_size=10, space_after=12)

    # === Delivery conditions ===
    _add_paragraph(doc, "Условия поставки:", bold=True, font_size=11, space_after=6)

    conditions = _build_delivery_conditions_list(data)
    for i, condition in enumerate(conditions, 1):
        # Handle multiline conditions (consignee)
        lines = condition.split('\n')
        for j, line in enumerate(lines):
            if j == 0:
                text = f"{i}. {line}"
            else:
                text = f"    {line}"  # Indent continuation lines
            _add_paragraph(doc, text, font_size=10, space_after=4)

    doc.add_paragraph()  # Spacer

    # === Footer text ===
    _add_paragraph(doc, "Настоящая Спецификация согласована и подписана Сторонами в 2 (двух) подлинных экземплярах.",
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=9, space_after=24)

    # === Signature blocks ===
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.columns[0].width = Cm(8)
    sig_table.columns[1].width = Cm(8)

    # Left column - Supplier
    left_cell = sig_table.cell(0, 0)
    left_p = left_cell.paragraphs[0]
    left_p.add_run("Поставщик").bold = True
    left_p.runs[0].font.size = Pt(10)

    left_cell.add_paragraph(seller_name).runs[0].font.size = Pt(10)
    left_cell.add_paragraph("Генеральный директор").runs[0].font.size = Pt(10)

    sig_p = left_cell.add_paragraph()
    sig_p.add_run(f"\n______________/ {seller_director} /").font.size = Pt(10)

    mp_p = left_cell.add_paragraph()
    mp_p.add_run("м.п.").font.size = Pt(8)

    # Right column - Customer
    right_cell = sig_table.cell(0, 1)
    right_p = right_cell.paragraphs[0]
    right_p.add_run("Покупатель").bold = True
    right_p.runs[0].font.size = Pt(10)

    right_cell.add_paragraph(customer_name).runs[0].font.size = Pt(10)
    right_cell.add_paragraph(signatory_position).runs[0].font.size = Pt(10)

    sig_p2 = right_cell.add_paragraph()
    sig_p2.add_run(f"\n______________/ {customer_signatory} /").font.size = Pt(10)

    mp_p2 = right_cell.add_paragraph()
    mp_p2.add_run("м.п.").font.size = Pt(8)

    # 5. Return DOCX bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
